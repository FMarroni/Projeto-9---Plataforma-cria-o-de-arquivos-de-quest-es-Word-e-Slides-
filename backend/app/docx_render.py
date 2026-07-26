"""Renderização de texto tokenizado (formatting.tokenize_rich_text) para
python-docx, incluindo injeção de imagens reais nos pontos marcados por
[IMAGEM_NN] e de fórmulas $$...$$ renderizadas (V3, ver latex_render.py).
Compartilhado por docx_lista.py e docx_comentada.py."""

import io

from docx.shared import Cm, RGBColor

from app.formatting import tokenize_rich_text
from app.latex_render import resolver_tokens_equacao

LARGURA_IMAGEM_CM = 12
LARGURA_EQUACAO_CM = 5  # fórmulas renderizadas são pequenas — não usar a largura de um gráfico


def inserir_paragrafo_tokenizado(
    ancora_paragraph,
    texto: str,
    imagens: dict[str, bytes],
    cor: RGBColor | None = None,
    negrito_forcado: bool = False,
):
    """Insere um ou mais parágrafos ANTES de `ancora_paragraph`, tokenizando
    `texto` (markdown-lite + LaTeX + [IMAGEM_NN]) e aplicando negrito/itálico/
    sublinhado/sobre-subscrito nativos do python-docx, além de embutir a
    imagem real quando encontra um placeholder de imagem ou uma fórmula
    $$...$$ (renderizada sob demanda e cacheada em `imagens`).

    `cor`/`negrito_forcado` sobrepõem a cor/negrito em TODOS os runs deste
    texto (usado pelo cabeçalho "N. (banca/orgao-ano) enunciado", que deve
    sair inteiro em roxo/negrito — as alternativas não usam esses parâmetros
    e saem na formatação normal)."""
    tokens = resolver_tokens_equacao(tokenize_rich_text(texto), imagens)
    # estilo "Normal" explícito (não deixado para inferência) — o parágrafo
    # imediatamente anterior a este pode ser um "Heading 1" (título de
    # matéria); sem isso, alguns leitores/versões do Word podem herdar
    # formatação do heading em vez de cair no padrão do documento.
    paragrafo_atual = ancora_paragraph.insert_paragraph_before("", style="Normal")

    for token in tokens:
        if token.new_paragraph:
            paragrafo_atual = ancora_paragraph.insert_paragraph_before("", style="Normal")
            continue

        if token.is_image:
            dados = imagens.get(token.image_id) if token.image_id else None
            if dados:
                largura = LARGURA_EQUACAO_CM if (token.image_id or "").startswith("EQUACAO_") else LARGURA_IMAGEM_CM
                run = paragrafo_atual.add_run()
                run.add_picture(io.BytesIO(dados), width=Cm(largura))
            continue

        if not token.text:
            continue

        run = paragrafo_atual.add_run(token.text)
        run.font.bold = (token.bold or negrito_forcado) or None
        run.font.italic = token.italic or None
        run.font.underline = token.underline or None
        if cor is not None:
            run.font.color.rgb = cor
        # superscript/subscript escrevem no MESMO elemento XML <w:vertAlign> —
        # setar os dois sempre (mesmo o que é False) faz o segundo apagar o
        # primeiro (Font.subscript=None remove o vertAlign inteiro). Só define
        # o que de fato está ativo.
        if token.superscript:
            run.font.superscript = True
        elif token.subscript:
            run.font.subscript = True

    return paragrafo_atual
