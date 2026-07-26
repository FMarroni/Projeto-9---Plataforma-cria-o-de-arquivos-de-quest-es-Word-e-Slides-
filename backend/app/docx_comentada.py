"""Gera o DOCX 'Questões Comentadas': enunciado + alternativas + imagens
injetadas, agrupadas por matéria; ao final de cada questão vem o comentário
(Coruj.IA) e só então o gabarito — nessa ordem, uma questão após a outra
(sem seção de gabarito consolidado separada)."""

import os

from docx import Document

from app.docx_common import COR_MARCA, cabecalho_inline, find_paragraph_with
from app.docx_render import inserir_paragrafo_tokenizado
from app.schemas import ExtractionResult


def _preencher_questoes_comentadas(
    doc: Document, extraction: ExtractionResult, imagens: dict[str, bytes]
) -> None:
    marcador = find_paragraph_with(doc, "{{QUESTOES_COMENTADAS}}")
    if marcador is None:
        return

    questoes_ordenadas = sorted(extraction.questoes, key=lambda q: (q.materia, q.numero))
    materia_atual = None

    for questao in questoes_ordenadas:
        if questao.materia != materia_atual:
            titulo_materia = marcador.insert_paragraph_before("", style="Heading 1")
            titulo_materia.add_run(questao.materia)
            materia_atual = questao.materia

        texto_enunciado = f"{questao.numero}. {cabecalho_inline(questao)} {questao.enunciado}"
        inserir_paragrafo_tokenizado(
            marcador, texto_enunciado, imagens, cor=COR_MARCA, negrito_forcado=True
        )

        for alt in questao.alternativas:
            inserir_paragrafo_tokenizado(marcador, f"{alt.letra}) {alt.texto}", imagens)

        if questao.comentario:
            comentario_par = marcador.insert_paragraph_before("")
            run_comentario = comentario_par.add_run("Comentários:")
            run_comentario.bold = True
            run_comentario.font.color.rgb = COR_MARCA
            inserir_paragrafo_tokenizado(marcador, questao.comentario, imagens)

        gabarito_texto = "Anulada" if questao.anulada else (questao.gabarito or "N/A")
        gabarito_par = marcador.insert_paragraph_before("")
        run_gabarito = gabarito_par.add_run(f"Gabarito: {gabarito_texto}")
        run_gabarito.bold = True
        run_gabarito.font.color.rgb = COR_MARCA

        marcador.insert_paragraph_before("")  # espaço entre questões

    marcador._p.getparent().remove(marcador._p)


def gerar_docx_comentada(
    extraction: ExtractionResult,
    template_path: str,
    saida_path: str,
    imagens: dict[str, bytes] | None = None,
) -> str:
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template não encontrado: {template_path}")

    imagens = imagens or {}
    doc = Document(template_path)

    _preencher_questoes_comentadas(doc, extraction, imagens)

    os.makedirs(os.path.dirname(saida_path), exist_ok=True)
    doc.save(saida_path)
    return saida_path
