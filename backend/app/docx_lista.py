"""Gera o DOCX 'Lista de Questões': só enunciado + alternativas + imagens
injetadas por questão, agrupadas por matéria; o gabarito de todas as questões
fica agrupado apenas no final do documento, em lista vertical (como o bloco
"Gabarito" do PDF original do TEC, um item por linha)."""

import os

from docx import Document

from app.docx_common import COR_MARCA, cabecalho_inline, find_paragraph_with
from app.docx_render import inserir_paragrafo_tokenizado
from app.schemas import ExtractionResult


def _gabarito_consolidado(extraction: ExtractionResult) -> str:
    questoes_ordenadas = sorted(extraction.questoes, key=lambda q: q.numero)
    linhas = [
        f"{q.numero}. {'Anulada' if q.anulada else (q.gabarito or 'N/A')}" for q in questoes_ordenadas
    ]
    return "\n".join(linhas)


def _preencher_lista_questoes(doc: Document, extraction: ExtractionResult, imagens: dict[str, bytes]) -> None:
    marcador = find_paragraph_with(doc, "{{LISTA_QUESTOES}}")
    if marcador is None:
        return

    questoes_ordenadas = sorted(extraction.questoes, key=lambda q: (q.materia, q.numero))
    materia_atual = None

    for questao in questoes_ordenadas:
        if questao.materia != materia_atual:
            titulo_materia = marcador.insert_paragraph_before("", style="Heading 1")
            titulo_materia.add_run(questao.materia)
            materia_atual = questao.materia

        texto_enunciado = f"{questao.numero}. {cabecalho_inline(questao)} {questao.enunciado}"
        inserir_paragrafo_tokenizado(
            marcador, texto_enunciado, imagens, cor=COR_MARCA, negrito_forcado=True
        )

        for alt in questao.alternativas:
            inserir_paragrafo_tokenizado(marcador, f"{alt.letra}) {alt.texto}", imagens)

        marcador.insert_paragraph_before("")  # espaço entre questões

    marcador._p.getparent().remove(marcador._p)


def gerar_docx_lista(
    extraction: ExtractionResult,
    template_path: str,
    saida_path: str,
    imagens: dict[str, bytes] | None = None,
) -> str:
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template não encontrado: {template_path}")

    imagens = imagens or {}
    doc = Document(template_path)

    _preencher_lista_questoes(doc, extraction, imagens)

    marcador_gabarito = find_paragraph_with(doc, "{{GABARITO_CONSOLIDADO}}")
    if marcador_gabarito is not None:
        for run in marcador_gabarito.runs:
            run.text = ""
        texto = _gabarito_consolidado(extraction)
        if marcador_gabarito.runs:
            marcador_gabarito.runs[0].text = texto
        else:
            marcador_gabarito.add_run(texto)

    os.makedirs(os.path.dirname(saida_path), exist_ok=True)
    doc.save(saida_path)
    return saida_path
