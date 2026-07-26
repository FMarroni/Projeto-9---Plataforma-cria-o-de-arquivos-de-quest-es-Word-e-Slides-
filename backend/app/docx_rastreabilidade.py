"""Gera o DOCX 'Rastreabilidade' (Épico 4 — Módulo Biblioteca): para cada
questão comentada em Modo Restrito (RAG), mostra o enunciado e, logo abaixo,
o arquivo+página do material de apoio que fundamentou cada alternativa/item
julgado — ou o aviso de "informação não encontrada", quando for o caso. Este
arquivo é uma trilha de auditoria para quem revisa o conteúdo; a citação de
fonte NUNCA aparece em comentada.docx/slides.pptx (ver docx_comentada.py/
pptx_gen.py) — o aluno final recebe só o comentário, não a fonte."""

import os

from docx import Document

from app.comments import MENSAGEM_INFO_NAO_ENCONTRADA
from app.docx_common import COR_MARCA, find_paragraph_with
from app.docx_render import inserir_paragrafo_tokenizado
from app.schemas import ExtractionResult, Questao

_AVISO_SEM_QUESTOES = "Nenhuma questão foi comentada em Modo Restrito (base de conhecimento) nesta análise."


def _foi_processada_em_modo_restrito(questao: Questao) -> bool:
    return bool(questao.rastreabilidade) or questao.comentario == MENSAGEM_INFO_NAO_ENCONTRADA


def _preencher_rastreabilidade(doc: Document, extraction: ExtractionResult) -> None:
    marcador = find_paragraph_with(doc, "{{RASTREABILIDADE}}")
    if marcador is None:
        return

    questoes = [
        q for q in sorted(extraction.questoes, key=lambda q: q.numero) if _foi_processada_em_modo_restrito(q)
    ]

    if not questoes:
        marcador.text = _AVISO_SEM_QUESTOES
        return

    for questao in questoes:
        titulo = marcador.insert_paragraph_before("", style="Heading 3")
        titulo.add_run(f"Questão {questao.numero}")

        inserir_paragrafo_tokenizado(marcador, questao.enunciado, {})

        if not questao.rastreabilidade:
            aviso_par = marcador.insert_paragraph_before("")
            aviso_par.add_run(MENSAGEM_INFO_NAO_ENCONTRADA)
        else:
            for item in questao.rastreabilidade:
                linha_par = marcador.insert_paragraph_before("")
                run = linha_par.add_run(
                    f"Alternativa {item.alternativa.upper()}: Fundamentada pelo arquivo "
                    f"{item.arquivo}, Página {item.pagina}."
                )
                run.font.color.rgb = COR_MARCA

        marcador.insert_paragraph_before("")  # espaço entre questões

    marcador._p.getparent().remove(marcador._p)


def gerar_docx_rastreabilidade(extraction: ExtractionResult, template_path: str, saida_path: str) -> str:
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template não encontrado: {template_path}")

    doc = Document(template_path)
    _preencher_rastreabilidade(doc, extraction)

    os.makedirs(os.path.dirname(saida_path), exist_ok=True)
    doc.save(saida_path)
    return saida_path
