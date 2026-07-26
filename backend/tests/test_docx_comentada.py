import os

from docx import Document

from app.docx_comentada import gerar_docx_comentada
from app.schemas import Alternativa, ExtractionResult, Questao

TEMPLATE_PATH = os.path.join(
    os.path.dirname(__file__), "..", "..", "templates", "comentada.docx"
)


def _extraction():
    return ExtractionResult(
        bancas=["FCC"],
        anos=[2025],
        cargos=["Analista"],
        questoes=[
            Questao(
                numero=1,
                banca="FCC",
                orgao="TRT 1",
                ano=2025,
                materia="Matemática",
                assunto="Assunto A",
                enunciado="Enunciado da questão 1.",
                alternativas=[Alternativa(letra="a", texto="alt a")],
                gabarito="B",
                comentario="Este é o comentário da questão 1.",
            ),
        ],
    )


def test_ordem_e_alternativas_depois_comentarios_depois_gabarito(tmp_path):
    saida = gerar_docx_comentada(_extraction(), TEMPLATE_PATH, str(tmp_path / "comentada.docx"))
    doc = Document(saida)
    textos = [p.text for p in doc.paragraphs if p.text.strip()]

    idx_alt = next(i for i, t in enumerate(textos) if t.startswith("a) alt a"))
    idx_comentarios = textos.index("Comentários:")
    idx_texto_comentario = next(
        i for i, t in enumerate(textos) if "Este é o comentário" in t
    )
    idx_gabarito = textos.index("Gabarito: B")

    assert idx_alt < idx_comentarios < idx_texto_comentario < idx_gabarito


def test_cabecalho_inline_no_enunciado(tmp_path):
    saida = gerar_docx_comentada(_extraction(), TEMPLATE_PATH, str(tmp_path / "comentada.docx"))
    doc = Document(saida)
    textos = [p.text for p in doc.paragraphs]
    assert "1. (FCC / TRT 1 - 2025) Enunciado da questão 1." in textos


def test_cabecalho_e_enunciado_saem_em_roxo_negrito_alternativa_nao(tmp_path):
    from docx.shared import RGBColor

    saida = gerar_docx_comentada(_extraction(), TEMPLATE_PATH, str(tmp_path / "comentada.docx"))
    doc = Document(saida)

    par_cabecalho = next(p for p in doc.paragraphs if p.text.startswith("1. (FCC"))
    assert all(r.font.bold for r in par_cabecalho.runs if r.text)
    assert all(r.font.color.rgb == RGBColor(0x42, 0x31, 0xA4) for r in par_cabecalho.runs if r.text)

    par_alternativa = next(p for p in doc.paragraphs if p.text == "a) alt a")
    assert all(not r.font.bold for r in par_alternativa.runs if r.text)
    assert all(r.font.color.rgb is None for r in par_alternativa.runs if r.text)


def test_titulo_de_materia_inserido(tmp_path):
    saida = gerar_docx_comentada(_extraction(), TEMPLATE_PATH, str(tmp_path / "comentada.docx"))
    doc = Document(saida)
    textos = [p.text for p in doc.paragraphs if p.text.strip()]
    assert "Matemática" in textos
