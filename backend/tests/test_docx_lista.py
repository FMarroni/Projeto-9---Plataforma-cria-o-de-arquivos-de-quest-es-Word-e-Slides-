import os

from docx import Document

from app.docx_lista import gerar_docx_lista
from app.formula_resolve import resolver_formulas_questao
from app.schemas import Alternativa, ExtractionResult, Formula, Questao

TEMPLATE_PATH = os.path.join(
    os.path.dirname(__file__), "..", "..", "templates", "lista.docx"
)


def _extraction():
    return ExtractionResult(
        bancas=["FCC"],
        anos=[2025],
        cargos=["Analista"],
        questoes=[
            Questao(
                numero=2,
                banca="FCC",
                orgao="TRT 1",
                ano=2025,
                materia="Direito",
                assunto="Assunto B",
                enunciado="Enunciado da questão 2.",
                alternativas=[Alternativa(letra="a", texto="alt a")],
                gabarito="A",
            ),
            Questao(
                numero=1,
                banca="FCC",
                orgao=None,
                ano=None,
                materia="Matemática",
                assunto="Assunto A",
                enunciado="Enunciado da questão 1.",
                alternativas=[Alternativa(letra="a", texto="alt a")],
                gabarito="B",
            ),
            Questao(
                numero=3,
                banca="FCC",
                orgao="TRT 1",
                ano=2024,
                materia="Direito",
                assunto="Assunto C",
                enunciado="Enunciado da questão 3.",
                alternativas=[],
                gabarito=None,
                anulada=True,
            ),
        ],
    )


def test_agrupa_por_materia_e_ordena_por_numero_dentro_do_grupo(tmp_path):
    saida = gerar_docx_lista(_extraction(), TEMPLATE_PATH, str(tmp_path / "lista.docx"))
    doc = Document(saida)
    textos = [p.text for p in doc.paragraphs if p.text.strip()]

    idx_direito = textos.index("Direito")
    idx_matematica = textos.index("Matemática")
    # "Direito" agrupa as questões 2 e 3 (nessa ordem); "Matemática" vem depois com a questão 1
    assert idx_direito < idx_matematica
    assert "2. (FCC / TRT 1 - 2025) Enunciado da questão 2." in textos
    assert "3. (FCC / TRT 1 - 2024) Enunciado da questão 3." in textos
    assert "1. (FCC) Enunciado da questão 1." in textos


def test_cabecalho_e_enunciado_saem_em_roxo_negrito_alternativa_nao(tmp_path):
    from docx.shared import RGBColor

    saida = gerar_docx_lista(_extraction(), TEMPLATE_PATH, str(tmp_path / "lista.docx"))
    doc = Document(saida)

    par_cabecalho = next(p for p in doc.paragraphs if p.text.startswith("2. (FCC"))
    assert all(r.font.bold for r in par_cabecalho.runs if r.text)
    assert all(r.font.color.rgb == RGBColor(0x42, 0x31, 0xA4) for r in par_cabecalho.runs if r.text)

    par_alternativa = next(p for p in doc.paragraphs if p.text == "a) alt a")
    assert all(not r.font.bold for r in par_alternativa.runs if r.text)
    assert all(r.font.color.rgb is None for r in par_alternativa.runs if r.text)


def test_nao_existe_mais_cabecalho_questao_n(tmp_path):
    saida = gerar_docx_lista(_extraction(), TEMPLATE_PATH, str(tmp_path / "lista.docx"))
    doc = Document(saida)
    textos = [p.text for p in doc.paragraphs]
    assert not any("Questão" in t and "—" in t for t in textos)


def test_gabarito_consolidado_e_vertical(tmp_path):
    saida = gerar_docx_lista(_extraction(), TEMPLATE_PATH, str(tmp_path / "lista.docx"))
    doc = Document(saida)

    marcador = next(p for p in doc.paragraphs if "1." in p.text and "2." in p.text)
    linhas = marcador.text.split("\n")
    assert linhas == ["1. B", "2. A", "3. Anulada"]


# --- regressão do bug relatado (G(t)=...) também no caminho DOCX ------------

LATEX_GT_CORRETO = r"G(t)=t^{3}-\frac{23}{2}t^{2}+\frac{55}{4}t+\frac{399}{8},\ t\in[0,10]"
LATEX_GT_ERRADO = r"G(t)=-\frac{23}{2}t^{2}+\frac{55}{4}t^{3}+\frac{399}{8}"


def test_gt_formula_correta_vira_imagem_no_docx_sem_texto_corrompido(tmp_path):
    questao = Questao(
        numero=1,
        banca="FCC",
        ano=2025,
        materia="Matemática",
        assunto="Funções",
        enunciado="Considere a função G definida por [FORMULA_01] no intervalo dado.",
        alternativas=[],
        gabarito="A",
        formulas=[Formula(id="FORMULA_01", latex=LATEX_GT_CORRETO, display=True, confidence=0.95)],
    )
    questao = resolver_formulas_questao(questao)
    assert "[FORMULA_01]" not in questao.enunciado

    extraction = ExtractionResult(bancas=["FCC"], anos=[2025], cargos=["Analista"], questoes=[questao])
    saida = gerar_docx_lista(extraction, TEMPLATE_PATH, str(tmp_path / "lista.docx"), {"FORMULA_01": b""})
    doc = Document(saida)

    # a fórmula foi renderizada e embutida como imagem inline -- não como
    # texto (que poderia estar corrompido/reordenado).
    assert len(doc.inline_shapes) >= 1
    for paragrafo in doc.paragraphs:
        assert LATEX_GT_ERRADO not in paragrafo.text
        assert r"\frac" not in paragrafo.text  # LaTeX bruto nunca sobra como texto visível
