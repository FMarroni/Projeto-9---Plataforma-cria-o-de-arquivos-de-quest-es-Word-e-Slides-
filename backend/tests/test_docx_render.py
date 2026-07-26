from docx import Document

from app.docx_render import inserir_paragrafo_tokenizado


def _runs_com_texto(doc: Document) -> list:
    return [r for p in doc.paragraphs for r in p.runs if r.text]


def test_sobrescrito_e_preservado_no_run():
    # regressão: Font.superscript e Font.subscript escrevem no MESMO elemento
    # XML <w:vertAlign> — setar subscript=None (via "or None") depois de
    # superscript=True apaga o vertAlign inteiro. Ver docx_render.py.
    doc = Document()
    marcador = doc.add_paragraph("MARCADOR")
    inserir_paragrafo_tokenizado(marcador, r"6x$^{2}$", {})

    runs = _runs_com_texto(doc)
    run_expoente = next(r for r in runs if r.text == "2")
    assert run_expoente.font.superscript is True
    assert run_expoente.font.subscript is not True


def test_subscrito_e_preservado_no_run():
    doc = Document()
    marcador = doc.add_paragraph("MARCADOR")
    inserir_paragrafo_tokenizado(marcador, r"H$_{2}$O", {})

    runs = _runs_com_texto(doc)
    run_indice = next(r for r in runs if r.text == "2")
    assert run_indice.font.subscript is True
    assert run_indice.font.superscript is not True


def test_texto_normal_nao_ganha_sobre_ou_subscrito():
    doc = Document()
    marcador = doc.add_paragraph("MARCADOR")
    inserir_paragrafo_tokenizado(marcador, "texto simples sem formatação", {})

    runs = _runs_com_texto(doc)
    assert all(r.font.superscript is not True for r in runs)
    assert all(r.font.subscript is not True for r in runs)


def test_moeda_dentro_do_paragrafo_nao_e_corrompida():
    doc = Document()
    marcador = doc.add_paragraph("MARCADOR")
    inserir_paragrafo_tokenizado(marcador, "Custou R$ 200,00 ao todo", {})

    texto_completo = "".join(r.text for r in _runs_com_texto(doc) if r.text != "MARCADOR")
    assert texto_completo == "Custou R$ 200,00 ao todo"
