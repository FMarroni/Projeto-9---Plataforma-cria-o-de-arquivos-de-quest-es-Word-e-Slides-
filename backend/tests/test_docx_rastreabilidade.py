import os

from docx import Document

from app.comments import MENSAGEM_INFO_NAO_ENCONTRADA
from app.docx_rastreabilidade import gerar_docx_rastreabilidade
from app.schemas import Alternativa, ExtractionResult, Questao, RastreabilidadeItem

TEMPLATE_PATH = os.path.join(
    os.path.dirname(__file__), "..", "..", "templates", "rastreabilidade.docx"
)


def _questao_com_rastreabilidade():
    return Questao(
        numero=1,
        banca="FCC",
        materia="Direito Penal",
        assunto="Furto",
        enunciado="O crime de furto exige subtração de coisa alheia móvel.",
        alternativas=[Alternativa(letra="a", texto="Certo"), Alternativa(letra="b", texto="Errado")],
        gabarito="a",
        comentario="(a) Correto, conforme o material de apoio.",
        rastreabilidade=[
            RastreabilidadeItem(alternativa="a", arquivo="Aula_01_Penal.pdf", pagina="14"),
            RastreabilidadeItem(alternativa="b", arquivo="Aula_01_Penal.pdf", pagina="16"),
        ],
    )


def _questao_sem_biblioteca():
    return Questao(
        numero=2,
        banca="FCC",
        materia="Direito Penal",
        assunto="Roubo",
        enunciado="Enunciado da questão 2.",
        gabarito="C",
        comentario="Comentário normal, sem RAG.",
    )


def _questao_info_nao_encontrada():
    return Questao(
        numero=3,
        banca="FCC",
        materia="Direito Penal",
        assunto="Estelionato",
        enunciado="Enunciado da questão 3.",
        gabarito="E",
        comentario=MENSAGEM_INFO_NAO_ENCONTRADA,
        rastreabilidade=[],
    )


def test_gera_secao_por_questao_com_rastreabilidade(tmp_path):
    extraction = ExtractionResult(questoes=[_questao_com_rastreabilidade()])
    saida = gerar_docx_rastreabilidade(extraction, TEMPLATE_PATH, str(tmp_path / "rastreabilidade.docx"))

    doc = Document(saida)
    textos = [p.text for p in doc.paragraphs if p.text.strip()]

    assert "Questão 1" in textos
    assert any("O crime de furto exige" in t for t in textos)
    assert "Alternativa A: Fundamentada pelo arquivo Aula_01_Penal.pdf, Página 14." in textos
    assert "Alternativa B: Fundamentada pelo arquivo Aula_01_Penal.pdf, Página 16." in textos


def test_questao_sem_biblioteca_e_omitida(tmp_path):
    extraction = ExtractionResult(questoes=[_questao_sem_biblioteca()])
    saida = gerar_docx_rastreabilidade(extraction, TEMPLATE_PATH, str(tmp_path / "rastreabilidade.docx"))

    doc = Document(saida)
    textos = [p.text for p in doc.paragraphs if p.text.strip()]

    assert "Questão 2" not in textos


def test_questao_info_nao_encontrada_mostra_aviso(tmp_path):
    extraction = ExtractionResult(questoes=[_questao_info_nao_encontrada()])
    saida = gerar_docx_rastreabilidade(extraction, TEMPLATE_PATH, str(tmp_path / "rastreabilidade.docx"))

    doc = Document(saida)
    textos = [p.text for p in doc.paragraphs if p.text.strip()]

    assert "Questão 3" in textos
    assert MENSAGEM_INFO_NAO_ENCONTRADA in textos


def test_sem_nenhuma_questao_em_modo_restrito_mostra_aviso_geral(tmp_path):
    extraction = ExtractionResult(questoes=[_questao_sem_biblioteca()])
    saida = gerar_docx_rastreabilidade(extraction, TEMPLATE_PATH, str(tmp_path / "rastreabilidade.docx"))

    doc = Document(saida)
    textos = [p.text for p in doc.paragraphs if p.text.strip()]

    assert any("Nenhuma questão foi comentada em Modo Restrito" in t for t in textos)


def test_mistura_questoes_com_e_sem_biblioteca(tmp_path):
    extraction = ExtractionResult(
        questoes=[_questao_com_rastreabilidade(), _questao_sem_biblioteca(), _questao_info_nao_encontrada()]
    )
    saida = gerar_docx_rastreabilidade(extraction, TEMPLATE_PATH, str(tmp_path / "rastreabilidade.docx"))

    doc = Document(saida)
    textos = [p.text for p in doc.paragraphs if p.text.strip()]

    assert "Questão 1" in textos
    assert "Questão 2" not in textos
    assert "Questão 3" in textos
